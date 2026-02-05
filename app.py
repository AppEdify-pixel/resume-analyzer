import streamlit as st
import re
import spacy
from io import BytesIO
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from fpdf import FPDF
import nltk
from nltk.tokenize import sent_tokenize
import unicodedata
from datetime import datetime
from collections import Counter

# 🔹 Minimal tech alias normalization (DO NOT expand arbitrarily)
TECH_ALIASES = {
    "asp.net": ["aspnet", "asp net"],
    ".net": ["net", "dotnet", "asp.net"],
    "c#": ["csharp"],
}

def normalize_with_aliases(text: str) -> str:
    """
    Normalizes known tech aliases so matching works
    without breaking punctuation like .NET or C#.
    """
    normalized = text.lower()
    for canonical, variants in TECH_ALIASES.items():
        for v in variants:
            normalized = re.sub(
                rf"\b{re.escape(v)}\b",
                canonical,
                normalized,
                flags=re.IGNORECASE,
            )
    return normalized

# added on Jan 27 2026
def alias_match_in_text(text, keyword):
    return normalize_with_aliases(keyword) in normalize_with_aliases(text)

# added lastly on Feb 03 2026
def highlight_dot_tech_token(para, text, keywords):
    """
    Highlights full dot-based tech tokens like ASP.NET or .NET
    as a single unit to avoid partial highlighting.
    """
    for kw in keywords:
        if "." in kw:
            if re.search(rf"(?<!\w){re.escape(kw)}(?!\w)", text, re.IGNORECASE):
            # 🔹 CLEAR ALL EXISTING RUNS FIRST added lastly on Feb 04 2026
               while para.runs:
                para._element.remove(para.runs[0]._element)

                run = para.add_run(text)
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                return True
    return False


# ----------------------------------------------------------------------
# ✅ Setup & Downloads
# ----------------------------------------------------------------------
for pkg in ["punkt", "punkt_tab"]:
    try:
        nltk.download(pkg, quiet=True)
    except:
        pass

# Initialize SpaCy
try:
    @st.cache_resource
def load_spacy_model():
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        import subprocess
        subprocess.run(
            ["python", "-m", "spacy", "download", "en_core_web_sm"],
            check=True
        )
        return spacy.load("en_core_web_sm")

nlp = load_spacy_model()

except:
    import os
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# ----------------------------------------------------------------------
# 🛠️ Helper Functions
# ----------------------------------------------------------------------

def normalize_text_for_pdf(text):
    """Removes/replaces Unicode characters that FPDF cannot encode."""
    if not text: return ""
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r'\s+', ' ', text).strip()
    replacements = {
        "—": "-", "–": "-", "―": "-", "…": "...", "•": "*",
        "‘": "'", "’": "'", "‚": "'", "‛": "'",
        "“": '"', "”": '"', "„": '"', "‟": '"',
        "′": "'", "″": '"', "‐": "-",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return ''.join(ch if ord(ch) < 256 else '?' for ch in text)

def iter_all_paragraphs(doc):
    """Yields all Paragraph objects in the document, including tables."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


    # def clean_jd_text(text):
    # """Adds spaces around symbols to prevent 'glued' words like PowerPlatform(PowerApps)."""
    # Force space around brackets, slashes, and colons
    # vtext = re.sub(r'([()/\\:])', r' \1 ', text)
    # return re.sub(r'\s+', ' ', text).strip()

def clean_jd_text(text):
    # Add space before/after dots only if not part of a known extension like .NET
    text = re.sub(r'(?<!\.)\.(?!\d|net|com|org)', ' . ', text, flags=re.IGNORECASE)
    # Force space around all common separators
    text = re.sub(r'([()/\\:;])', r' \1 ', text)
    return re.sub(r'\s+', ' ', text).strip()

def get_proper_case(word):
    """Handles consistent casing for both the Resume and PDF Summary."""
    canonical = {
        "sql": "SQL", "m365": "M365", "api": "API", "aws": "AWS", 
        "ci/cd": "CI/CD", "c#": "C#", ".net": ".NET", "mvc": "MVC",
        "id": "ID", "ad": "AD", "bi": "BI"
    }
    low = word.lower()
    if low in canonical:
        return canonical[low]
    # Acronyms (3 letters or less) become UPPERCASE, others become Title Case
    return low.upper() if len(low) <= 3 else low.title()

# ----------------------------------------------------------------------
# 🧠 1️⃣ Keyword Extraction
# ----------------------------------------------------------------------
def extract_keywords(text):
    if not text: return []
    doc = nlp(text)
    potential_keywords = set()
    
    # "Security Guard" list to block non-tech header words
    exclusion_list = {
        "candidate", "requirement", "requirements", "position", "positions", "skills", 
        "additional", "summary", "experience", "qualifications", "objective", "profile",
        "education", "responsibilities", "deliverables", "management", "project", 
        "business", "delivery", "client", "tools", "solutions", "system", "systems",
        "information", "description", "details", "contact", "level", "years", "minimum",
        "preferred", "standard", "procedures", "process", "various", "multiple",
        "mandatory", "technical", "tasks", "skills/qualification", "responsibilities",
        "suite", "related", "organizational","the", "based", "related", 
        "entire", "various", "use", "including"
    }

    clean_pattern = r"[^\w\s\-\#\+]"

    # A. Noun Chunks (Phrases)
    for chunk in doc.noun_chunks:
        if any(token.pos_ == "PROPN" or token.text.isupper() for token in chunk):
            cleaned_phrase = re.sub(clean_pattern, '', chunk.text).strip()
            words = cleaned_phrase.lower().split()
            # Only pick phrases of 2-3 words to avoid long sentences
            if 1 < len(words) <= 3 and not any(w in exclusion_list for w in words):
                potential_keywords.add(cleaned_phrase.lower())

    # B. Single Words (C#, SQL, Microsoft)
    for token in doc:
        text_clean = re.sub(clean_pattern, '', token.text).strip()
        is_acronym = token.text.isupper() and len(text_clean) > 1
        is_special_tech = text_clean.lower() in ['c#', 'c++', '.net']
        is_proper_name = token.pos_ == "PROPN" and len(text_clean) > 2
        
        if (is_acronym or is_special_tech or is_proper_name):
            if text_clean.lower() not in exclusion_list:
                potential_keywords.add(text_clean.lower())

    return sorted(list(potential_keywords))

# ----------------------------------------------------------------------
# 🧩 2️⃣ Highlight + Capitalize
# ----------------------------------------------------------------------
def highlight_and_capitalize_docx(doc, keywords):
    lower_keywords = [k.lower() for k in keywords]
    # Use (?!\w) instead of \b to correctly identify C# and .NET
    pattern = r'(\b' + r'(?!\w)|\b'.join([re.escape(k) for k in lower_keywords]) + r'(?!\w))'
    
    fix_count = 0
    cap_changes = []

    for para in iter_all_paragraphs(doc): 
        original_text = para.text 
        if not original_text.strip(): continue
        
        # 🔹 FIX: handle ASP.NET / .NET as a single token
        if highlight_dot_tech_token(para, original_text, lower_keywords):
            # while len(para.runs) > 1:
            #    para._element.remove(para.runs[0]._element)
            continue
            
        parts = re.split(pattern, original_text, flags=re.IGNORECASE)
        while len(para.runs) > 0:
            para._element.remove(para.runs[0]._element)
        
        for part in parts:
            if not part: continue
            new_run = para.add_run(part)
            if normalize_with_aliases(part) in normalize_with_aliases(" ".join(lower_keywords)):
                proper_case = get_proper_case(part)
                new_run.text = proper_case
                new_run.bold = True
                new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                if part != proper_case:
                    cap_changes.append((part, proper_case))
                    fix_count += 1
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), fix_count, cap_changes

# ----------------------------------------------------------------------
# 🧾 3️⃣ PDF Summary Report
# ----------------------------------------------------------------------
def generate_summary_pdf(found, missing, cap_fixes, capitalization_changes):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Resume Analysis Summary", ln=True, align="C")
    pdf.ln(6)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 8, f"Generated on: {now}", ln=True)
    pdf.ln(4)

    # Capitalization Section
    grouped = Counter([after for _, after in capitalization_changes])
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, f"Capitalization Fixes Applied: {cap_fixes}", ln=True)
    pdf.set_font("Arial", size=10)
    for word, count in grouped.items():
        pdf.multi_cell(0, 5, f"- {normalize_text_for_pdf(word)}: {count} time{'s' if count > 1 else ''}", 0, 1)

    pdf.ln(8)
    
    # Matched Keywords
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, f"Matched Keywords ({len(found)}):", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 5, ", ".join([f"[OK] {get_proper_case(f)}" for f in found]), 0, 1)

    pdf.ln(5)
    
    # Missing Keywords
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, f"Missing Keywords ({len(missing)}):", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 5, ", ".join([f"[MISSING] {get_proper_case(m)}" for m in missing]), 0, 1)

    pdf_bytes = pdf.output(dest='S').encode('latin-1', errors='ignore')
    return pdf_bytes

# ----------------------------------------------------------------------
# 🖥️ 4️⃣ Streamlit Interface
# ----------------------------------------------------------------------
st.set_page_config(page_title="Resume Analyzer", page_icon="💼")
st.title("💼 Resume Keyword & Capitalization Analyzer")

req_file = st.file_uploader("📄 Upload Requirement Document", type=["pdf", "docx"])
res_file = st.file_uploader("👤 Upload Resume Document", type=["docx"])

if req_file and res_file and st.button("🔍 Analyze Resume"):
    with st.spinner("Processing..."):
        # Text Extraction Logic
        def extract_text_from_docx(file):
            d = Document(file)
            return " ".join(p.text for p in d.paragraphs)

        def extract_text_from_pdf(file):
            from PyPDF2 import PdfReader
            reader = PdfReader(file)
            return " ".join(page.extract_text() or "" for page in reader.pages)

        # 1. Extract and Clean Requirement Text
        if req_file.name.endswith(".pdf"):
            req_text = extract_text_from_pdf(req_file)
        else:
            req_text = extract_text_from_docx(req_file)
        
        req_text = clean_jd_text(req_text) # Fix "glued" words

        # 2. Extract Keywords
        keywords = extract_keywords(req_text)

        # 3. Process Resume
        resume_doc = Document(res_file)
        highlighted_bytes, cap_fixes, capitalization_changes = highlight_and_capitalize_docx(resume_doc, keywords)
        
        # Get flattened resume text for "Found" check
        resume_text = " ".join(p.text for p in iter_all_paragraphs(Document(res_file)))
        
        def phrase_components_present(phrase, resume_text):
            """
            Returns True if ALL meaningful components of a phrase
            are present somewhere in the resume text.
            """
            # Words to ignore when matching phrases
            weak_words = {
                "and", "or", "the", "of", "for", "with",
                "product", "products", "system", "systems",
                "solution", "solutions", "platform", "platforms",
                "enhancement", "enhancements"
            }
        
            components = [
                w for w in re.findall(r"[A-Za-z0-9]+", phrase.lower())
                if w not in weak_words
            ]
        
            return all(
                re.search(rf"\b{re.escape(c)}\b", resume_text, re.IGNORECASE)
                or normalize_with_aliases(c) in normalize_with_aliases(resume_text)
                for c in components
            )

        
        
        found = []
        missing = []
        
        normalized_resume_text = normalize_with_aliases(resume_text)
        
        
        for k in keywords:
            normalized_k = normalize_with_aliases(k)
            if phrase_components_present(normalized_k, normalized_resume_text):
                found.append(k)
            else:
                missing.append(k)


        # 4. Sync "Found" logic with Highlighter Regex
        # found = []
        # missing = []
        # for k in keywords:
        #    if re.search(rf"\b{re.escape(k)}(?!\w)", resume_text, re.IGNORECASE):
        #        found.append(k)
        #    else:
        #        missing.append(k)

        summary_pdf = generate_summary_pdf(found, missing, cap_fixes, capitalization_changes)

        st.session_state.update({
            "found": found, "missing": missing, "cap_fixes": cap_fixes,
            "highlighted_bytes": highlighted_bytes, "summary_pdf": summary_pdf, "analyzed": True
        })
        st.success("✅ Analysis Complete!")

# ----------------------------------------------------------------------
# 📊 5️⃣ Results Display
# ----------------------------------------------------------------------
if st.session_state.get("analyzed", False):
    st.subheader("📊 Summary Report")
    f_disp = [get_proper_case(k) for k in st.session_state["found"]]
    m_disp = [get_proper_case(k) for k in st.session_state["missing"]]

    st.markdown(f"**Matched ({len(f_disp)}):** " + ", ".join(f_disp))
    st.markdown(f"**Missing ({len(m_disp)}):** " + ", ".join(m_disp))

    col1, col2 = st.columns(2)
    with col1:
        st.download_button("📥 Corrected Resume (.docx)", data=st.session_state["highlighted_bytes"], file_name="Corrected_Resume.docx")
    with col2:

        st.download_button("📄 Summary Report (.pdf)", data=st.session_state["summary_pdf"], file_name="Summary_Report.pdf")
